"""Shared Lloyd's Register .docx parser.

Sprint D6.93 — Lloyd's Register publishes their Rules through the
Regs4ships portal in .docx form (one file per chapter, plus a General
Regulations doc and per-edition Notice docs that summarize amendments).

This module provides the shared parsing logic used by two downstream
source adapters:

  * ``lr_lifting_code`` — LR-CO-001 (Code for Lifting Appliances in a
    Marine Environment). Covers cranes, derricks, ro-ro access, lifts,
    shiplifts, cargo gear. Used by every vessel with lifting equipment.

  * ``lr_rules`` — LR-RU-001 (Rules and Regulations for the
    Classification of Ships). The big one. Covers hull construction,
    machinery, electrical engineering, surveys, ship types. This is
    the document Karynn's transformer-failure question needs.

Both share the same .docx structure (Chapter X.docx with Sections
inside, plus a General Regulations / Notice files) so one parser
handles both via the ``source_name`` + ``doc_prefix`` parameters.

Chunking strategy: each top-level Section becomes one Section object.
Granular sub-sections (e.g. "1.1.1", "2.3.5") stay inside the parent
section's full_text — they're typically short and lose context if
split apart. The result is ~10-30 chunks per chapter, which sits in
the same size range as our SOLAS / MARPOL / STCW chunks.
"""
from __future__ import annotations

import logging
import re
from datetime import date
from pathlib import Path

from ingest.models import Section

logger = logging.getLogger(__name__)

# Matches the explicit "Section N <Title>" headers Lloyd's uses inside
# every chapter (e.g. "Section 1 Introduction", "Section 2 Control,
# alarm and safety systems"). Used to split a chapter into Section
# objects.
_SECTION_HEADER = re.compile(
    r"^\s*Section\s+(\d+)(?:\s+(.+?))?\s*$",
    re.IGNORECASE,
)

# Matches the "Chapter N" pattern inside the filename. We use the
# filename rather than in-doc text because the .docx exports often
# strip the chapter banner from page 1.
_CHAPTER_FROM_FILENAME = re.compile(
    r"^Chapter\s+(\d+)\s+(.+?)\.docx$",
    re.IGNORECASE,
)


def _read_docx_paragraphs(docx_path: Path) -> list[str]:
    """Extract non-empty paragraphs from a .docx file in document order.

    Uses python-docx; tables are captured by walking each table's cell
    paragraphs after the body run. The order is approximate (python-docx
    doesn't expose interleaved table positions in body order without
    walking the underlying XML), which is acceptable for our chunking
    granularity since chapter/section boundaries are always paragraph-
    level headers, not table contents.
    """
    from docx import Document  # local import to keep module import light

    doc = Document(str(docx_path))
    paragraphs: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            paragraphs.append(text)
    # Append table cell text after the body. Loses ordering but
    # preserves content (LR puts examination matrices in tables).
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = (p.text or "").strip()
                    if text:
                        paragraphs.append(text)
    return paragraphs


def _split_into_sections(paragraphs: list[str]) -> list[tuple[str, str, str]]:
    """Split a flat paragraph stream into (section_num, title, body) tuples.

    Splits on lines matching "Section N <Title>". Anything before the
    first Section header is collected as a preamble keyed as
    section_num="0". If no sections are found, returns the entire body
    as a single preamble.
    """
    splits: list[tuple[int, str, list[str]]] = []  # (sec_num, title, body)
    current_num = 0
    current_title = "Preamble"
    current_body: list[str] = []

    for p in paragraphs:
        m = _SECTION_HEADER.match(p)
        if m:
            # Save the previous section's body, start a new one.
            if current_body or splits:
                splits.append((current_num, current_title, current_body))
            current_num = int(m.group(1))
            current_title = (m.group(2) or f"Section {current_num}").strip()
            current_body = []
        else:
            current_body.append(p)
    # Tail
    if current_body or splits:
        splits.append((current_num, current_title, current_body))

    return [(str(num), title, "\n\n".join(body).strip()) for num, title, body in splits]


def _detect_chapter(docx_path: Path) -> tuple[str, str]:
    """Return (chapter_number_str, chapter_title) for the given file.

    Examples:
      "Chapter 10 Electrotechnical Systems.docx" → ("10", "Electrotechnical Systems")
      "General Regulations.docx" → ("GenReg", "General Regulations")
      "Notice No.1 Code for Lifting...docx" → ("Notice1", "<full notice title>")
    """
    name = docx_path.name
    m = _CHAPTER_FROM_FILENAME.match(name)
    if m:
        return m.group(1), m.group(2).strip()
    # Notice file convention
    notice_m = re.match(r"Notice\s+No\.?\s*(\d+)\s+(.+?)\.docx$", name, re.IGNORECASE)
    if notice_m:
        return f"Notice{notice_m.group(1)}", notice_m.group(2).strip()
    # General Regulations
    if name.lower().startswith("general regulations"):
        return "GenReg", "General Regulations"
    # Fallback: use stem
    return docx_path.stem, docx_path.stem


def parse_lloyds_docx_dir(
    raw_dir: Path,
    *,
    source_name: str,
    doc_prefix: str,
    source_date: date,
) -> list[Section]:
    """Parse every .docx in ``raw_dir`` into Section objects.

    Args:
        raw_dir: Directory containing the .docx files (one per chapter
                 plus auxiliary docs).
        source_name: DB source value, e.g. "lr_lifting_code".
        doc_prefix: Citation prefix mariners would use, e.g. "LR-CO-001"
                    for the Code for Lifting Appliances, "LR-RU-001"
                    for the Rules for Classification of Ships.
        source_date: When this corpus was published (drives the
                     up_to_date_as_of column).
    """
    if not raw_dir.exists():
        raise FileNotFoundError(f"Lloyd's raw directory not found: {raw_dir}")

    docx_files = sorted(raw_dir.glob("*.docx"))
    if not docx_files:
        logger.warning("No .docx files found in %s", raw_dir)
        return []

    sections: list[Section] = []
    for docx_path in docx_files:
        chapter_num, chapter_title = _detect_chapter(docx_path)
        paragraphs = _read_docx_paragraphs(docx_path)
        if not paragraphs:
            logger.warning("Empty document: %s", docx_path.name)
            continue

        chapter_label = (
            f"Ch.{chapter_num}" if chapter_num.isdigit() else chapter_num
        )

        splits = _split_into_sections(paragraphs)
        for sec_num, sec_title, body in splits:
            if not body:
                continue
            # Section number convention:
            #   "LR-CO-001 Ch.10 Sec.2 — Control, alarm and safety systems"
            # mariners will write either "LR-CO-001 Ch.10 Sec.2" or
            # "Lloyd's CO-001 Ch.10 Sec.2"; the citation regex on the
            # frontend bridges that.
            if sec_num == "0":
                # Preamble — no section number, use chapter as anchor.
                section_number = f"{doc_prefix} {chapter_label}"
                section_title = f"{chapter_title}"
            else:
                section_number = f"{doc_prefix} {chapter_label} Sec.{sec_num}"
                section_title = f"{chapter_title} — {sec_title}"

            sections.append(Section(
                source=source_name,
                title_number=0,
                section_number=section_number,
                section_title=section_title[:500],  # DB column safety
                full_text=body,
                up_to_date_as_of=source_date,
                parent_section_number=f"{doc_prefix} {chapter_label}",
            ))

    logger.info(
        "Lloyd's %s: %d sections parsed across %d files",
        source_name, len(sections), len(docx_files),
    )
    return sections


def dry_run_dir(
    raw_dir: Path,
    *,
    source_name: str,
    doc_prefix: str,
    source_date: date,
) -> None:
    """Print parsed section summary without DB writes."""
    sections = parse_lloyds_docx_dir(
        raw_dir,
        source_name=source_name,
        doc_prefix=doc_prefix,
        source_date=source_date,
    )
    print(f"\n{source_name}: {len(sections)} sections parsed\n")
    for s in sections[:30]:
        preview = s.full_text[:80].replace("\n", " ")
        print(f"  [{s.section_number}]")
        print(f"    {s.section_title}")
        print(f"    {len(s.full_text):,} chars | {preview}...")
        print()
    if len(sections) > 30:
        print(f"  ... and {len(sections) - 30} more sections")
